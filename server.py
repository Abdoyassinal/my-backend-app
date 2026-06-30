from fastapi import FastAPI, APIRouter, HTTPException, Header
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime, timezone

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

APP_PASSWORD = "Alyassin-1997"
APP_TOKEN = "almahmod-2026-secret-token"

app = FastAPI()
api_router = APIRouter(prefix="/api")


# ============ MODELS ============
class LoginRequest(BaseModel):
    password: str


class Settings(BaseModel):
    company_name: str = "ALMAHMOD HANDEL"
    sender_info: str = "ALMAHMOD SAMIRA\nEichendorff Str.1\n31311 UETZE\n+49 17622314161\naleasen330@gmail.com\nSt-Nr: 16/101/16353\nSt-IdNr: DE458573396"
    footer: str = "ALMAHMOD HANDEL · Sparkasse Wolfsburg · BLZ 269 513 11 · Konto 0163539646 · IBAN DE44 2695 1311 0163 5396 46 · BIC NOLADE21GFW"
    tax_rate: float = 19.0


class CatalogItem(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name_de: str
    name_ar: str = ""
    default_price: float = 0.0


class CatalogItemCreate(BaseModel):
    name_de: str
    name_ar: str = ""
    default_price: float = 0.0


class InvoiceLine(BaseModel):
    description: str
    unit_price: float
    quantity: float


class Invoice(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    invoice_number: str
    invoice_date: str
    delivery_date: str
    recipient: str
    payment_method: str = "Bar Bezahlt"
    items: List[InvoiceLine]
    tax_rate: float = 19.0
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


class InvoiceCreate(BaseModel):
    invoice_number: str
    invoice_date: str
    delivery_date: str
    recipient: str
    payment_method: str = "Bar Bezahlt"
    items: List[InvoiceLine]
    tax_rate: float = 19.0


# ============ AUTH HELPER ============
def check_auth(authorization: Optional[str]):
    if not authorization or authorization != f"Bearer {APP_TOKEN}":
        raise HTTPException(status_code=401, detail="Unauthorized")


# ============ AUTH ENDPOINT ============
@api_router.post("/auth/login")
async def login(data: LoginRequest):
    if data.password != APP_PASSWORD:
        raise HTTPException(status_code=401, detail="كلمة السر غير صحيحة")
    return {"token": APP_TOKEN}


# ============ SETTINGS ============
DEFAULT_SETTINGS = Settings().model_dump()


@api_router.get("/settings")
async def get_settings(authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    doc = await db.settings.find_one({"_id": "default"}, {"_id": 0})
    if not doc:
        await db.settings.insert_one({"_id": "default", **DEFAULT_SETTINGS})
        return DEFAULT_SETTINGS
    return doc


@api_router.put("/settings")
async def update_settings(data: Settings, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    await db.settings.update_one(
        {"_id": "default"},
        {"$set": data.model_dump()},
        upsert=True,
    )
    return data


# ============ CATALOG ============
DEFAULT_CATALOG = [
    {"name_ar": "ملابس صلاة", "name_de": "Gebetskleidung", "default_price": 12.0},
    {"name_ar": "مفرش طاولة", "name_de": "Tischdecke", "default_price": 3.0},
    {"name_ar": "ملعقة", "name_de": "Löffel", "default_price": 1.5},
    {"name_ar": "شوكة", "name_de": "Gabeln", "default_price": 1.5},
    {"name_ar": "سيخ شواء", "name_de": "Grillspieß", "default_price": 1.5},
    {"name_ar": "قدور فخار", "name_de": "Töpfe aus Ton", "default_price": 12.0},
    {"name_ar": "سكين", "name_de": "Messer", "default_price": 2.0},
    {"name_ar": "صحن", "name_de": "Teller", "default_price": 2.5},
    {"name_ar": "كوب", "name_de": "Tasse", "default_price": 1.5},
    {"name_ar": "إبريق", "name_de": "Kanne", "default_price": 5.0},
]


async def ensure_catalog_seed():
    count = await db.catalog.count_documents({})
    if count == 0:
        for item in DEFAULT_CATALOG:
            obj = CatalogItem(**item)
            await db.catalog.insert_one({"_id": obj.id, **obj.model_dump()})


@api_router.get("/catalog")
async def list_catalog(authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    await ensure_catalog_seed()
    items = await db.catalog.find({}, {"_id": 0}).to_list(1000)
    return items


@api_router.post("/catalog")
async def create_catalog_item(data: CatalogItemCreate, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    obj = CatalogItem(**data.model_dump())
    await db.catalog.insert_one({"_id": obj.id, **obj.model_dump()})
    return obj


@api_router.put("/catalog/{item_id}")
async def update_catalog_item(item_id: str, data: CatalogItemCreate, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    result = await db.catalog.update_one(
        {"_id": item_id},
        {"$set": data.model_dump()},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Item not found")
    doc = await db.catalog.find_one({"_id": item_id}, {"_id": 0})
    return doc


@api_router.delete("/catalog/{item_id}")
async def delete_catalog_item(item_id: str, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    await db.catalog.delete_one({"_id": item_id})
    return {"ok": True}


# ============ INVOICES ============
@api_router.get("/invoices/last-number")
async def get_last_invoice_number(authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    doc = await db.invoices.find_one(
        {},
        sort=[("created_at", -1)],
        projection={"_id": 0, "invoice_number": 1, "created_at": 1},
    )
    return {"last_invoice_number": doc["invoice_number"] if doc else None}


@api_router.get("/invoices/stats")
async def invoice_stats(authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    docs = await db.invoices.find({}, {"_id": 0}).to_list(2000)
    total_count = len(docs)
    total_revenue = 0.0
    month_count = 0
    month_revenue = 0.0
    now = datetime.now(timezone.utc)
    this_month_prefix = f"{now.year}-{now.month:02d}"
    for d in docs:
        net = sum(i["unit_price"] * i["quantity"] for i in d.get("items", []))
        gross = net * (1 + d.get("tax_rate", 19) / 100)
        total_revenue += gross
        if (d.get("created_at") or "").startswith(this_month_prefix):
            month_count += 1
            month_revenue += gross
    return {
        "total_count": total_count,
        "total_revenue": round(total_revenue, 2),
        "month_count": month_count,
        "month_revenue": round(month_revenue, 2),
    }


@api_router.get("/invoices")
async def list_invoices(authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    docs = await db.invoices.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return docs


@api_router.get("/invoices/{invoice_id}")
async def get_invoice(invoice_id: str, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    doc = await db.invoices.find_one({"_id": invoice_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Invoice not found")
    return doc


@api_router.post("/invoices")
async def create_invoice(data: InvoiceCreate, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    obj = Invoice(**data.model_dump())
    await db.invoices.insert_one({"_id": obj.id, **obj.model_dump()})
    return obj


@api_router.put("/invoices/{invoice_id}")
async def update_invoice(invoice_id: str, data: InvoiceCreate, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    existing = await db.invoices.find_one({"_id": invoice_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Invoice not found")
    update_data = data.model_dump()
    update_data["id"] = invoice_id
    update_data["created_at"] = existing.get("created_at")
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.invoices.update_one({"_id": invoice_id}, {"$set": update_data})
    doc = await db.invoices.find_one({"_id": invoice_id}, {"_id": 0})
    return doc


@api_router.delete("/invoices/{invoice_id}")
async def delete_invoice(invoice_id: str, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    await db.invoices.delete_one({"_id": invoice_id})
    return {"ok": True}


# ============ EXPORT HELPERS ============
async def get_full_settings():
    doc = await db.settings.find_one({"_id": "default"}, {"_id": 0})
    return doc or DEFAULT_SETTINGS


def calculate_totals(items: List[InvoiceLine], tax_rate: float):
    net = sum(line.unit_price * line.quantity for line in items)
    tax = round(net * tax_rate / 100, 2)
    total = round(net + tax, 2)
    return round(net, 2), tax, total


def build_pdf(invoice: Invoice, settings: dict) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
    )
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=20, alignment=TA_CENTER, spaceAfter=20, textColor=colors.black, fontName='Helvetica-Bold')
    elements.append(Paragraph(settings.get("company_name", "ALMAHMOD HANDEL"), title_style))

    normal = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=9, leading=12)
    bold = ParagraphStyle('Bold', parent=styles['Normal'], fontSize=9, leading=12, fontName='Helvetica-Bold')

    # Top header: AN (left) and VON (right)
    recipient_html = "<b>AN</b><br/>" + invoice.recipient.replace("\n", "<br/>")
    sender_html = "<b>VON</b><br/>" + settings.get("sender_info", "").replace("\n", "<br/>")
    header_table = Table([[Paragraph(recipient_html, normal), Paragraph(sender_html, normal)]], colWidths=[8.5*cm, 8.5*cm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 0.8*cm))

    rechnung_style = ParagraphStyle('Rechnung', parent=styles['Heading2'], fontSize=16, textColor=colors.HexColor("#16a34a"), fontName='Helvetica-Bold', spaceAfter=10)
    elements.append(Paragraph("Rechnung", rechnung_style))

    # Invoice info table
    info_data = [
        ["Rechnungsnummer", invoice.invoice_number],
        ["Rechnungsdatum", invoice.invoice_date],
        ["Lieferdatum", invoice.delivery_date],
    ]
    info_table = Table(info_data, colWidths=[6*cm, 11*cm])
    info_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.5*cm))

    # Items table
    item_rows = [["Pos", "Leistung", "Einzelpreis", "Anzahl", "Gesamtpreis"]]
    for i, line in enumerate(invoice.items, start=1):
        item_rows.append([
            str(i),
            line.description,
            f"{line.unit_price:.2f} €",
            f"{line.quantity:g}",
            f"{line.unit_price * line.quantity:.2f} €",
        ])
    # Add empty rows to match invoice template style
    min_rows = 8
    while len(item_rows) - 1 < min_rows:
        item_rows.append(["", "", "", "", ""])

    items_table = Table(item_rows, colWidths=[1.2*cm, 8.3*cm, 2.5*cm, 2*cm, 3*cm])
    items_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f3f4f6")),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(items_table)
    elements.append(Spacer(1, 0.3*cm))

    # Totals
    net, tax, total = calculate_totals(invoice.items, invoice.tax_rate)
    totals_data = [
        ["NETTOBETRAG", f"{net:.2f} €"],
        [f"MWST {invoice.tax_rate:g}%", f"{tax:.2f} €"],
        [f"RECHNUNGSBETRAG ( {invoice.payment_method} )", f"{total:.2f} €"],
    ]
    totals_table = Table(totals_data, colWidths=[12*cm, 5*cm])
    totals_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor("#dcfce7")),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    elements.append(totals_table)

    # Footer
    footer_text = settings.get("footer", "")
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER, textColor=colors.HexColor("#374151"))

    def add_footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont('Helvetica', 7)
        canvas.setFillColor(colors.HexColor("#374151"))
        canvas.drawCentredString(A4[0] / 2, 1*cm, footer_text)
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)
    return buffer.getvalue()


def set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def build_docx(invoice: Invoice, settings: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(settings.get("company_name", "ALMAHMOD HANDEL"))
    run.bold = True
    run.font.size = Pt(20)

    # AN / VON header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    left_cell.text = ""
    p1 = left_cell.paragraphs[0]
    r = p1.add_run("AN\n")
    r.bold = True
    p1.add_run(invoice.recipient)

    right_cell.text = ""
    p2 = right_cell.paragraphs[0]
    r = p2.add_run("VON\n")
    r.bold = True
    p2.add_run(settings.get("sender_info", ""))

    doc.add_paragraph("")

    rech = doc.add_paragraph()
    r = rech.add_run("Rechnung")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

    # Info table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("Rechnungsnummer", invoice.invoice_number),
        ("Rechnungsdatum", invoice.invoice_date),
        ("Lieferdatum", invoice.delivery_date),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v

    doc.add_paragraph("")

    # Items table
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    hdr = items_table.rows[0].cells
    headers = ["Pos", "Leistung", "Einzelpreis", "Anzahl", "Gesamtpreis"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        set_cell_bg(hdr[i], "F3F4F6")

    for i, line in enumerate(invoice.items, start=1):
        row = items_table.add_row().cells
        row[0].text = str(i)
        row[1].text = line.description
        row[2].text = f"{line.unit_price:.2f} €"
        row[3].text = f"{line.quantity:g}"
        row[4].text = f"{line.unit_price * line.quantity:.2f} €"

    # Add empty rows
    min_rows = 8
    while len(items_table.rows) - 1 < min_rows:
        items_table.add_row()

    doc.add_paragraph("")

    # Totals
    net, tax, total = calculate_totals(invoice.items, invoice.tax_rate)
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Table Grid'
    totals = [
        ("NETTOBETRAG", f"{net:.2f} €", None),
        (f"MWST {invoice.tax_rate:g}%", f"{tax:.2f} €", None),
        (f"RECHNUNGSBETRAG ( {invoice.payment_method} )", f"{total:.2f} €", "DCFCE7"),
    ]
    for i, (k, v, bg) in enumerate(totals):
        cells = totals_table.rows[i].cells
        cells[0].text = ""
        cells[1].text = ""
        p1 = cells[0].paragraphs[0]
        run = p1.add_run(k)
        run.bold = True
        p2 = cells[1].paragraphs[0]
        run = p2.add_run(v)
        run.bold = True
        if bg:
            set_cell_bg(cells[0], bg)
            set_cell_bg(cells[1], bg)

    # Footer in section
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(settings.get("footer", ""))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@api_router.post("/invoices/export/pdf")
async def export_pdf(data: InvoiceCreate, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    settings = await get_full_settings()
    invoice = Invoice(**data.model_dump())
    pdf_bytes = build_pdf(invoice, settings)
    filename = f"Rechnung_{invoice.invoice_number}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.post("/invoices/export/docx")
async def export_docx(data: InvoiceCreate, authorization: Optional[str] = Header(None)):
    check_auth(authorization)
    settings = await get_full_settings()
    invoice = Invoice(**data.model_dump())
    docx_bytes = build_docx(invoice, settings)
    filename = f"Rechnung_{invoice.invoice_number}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api_router.get("/")
async def root():
    return {"message": "ALMAHMOD HANDEL - Invoice API"}


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
